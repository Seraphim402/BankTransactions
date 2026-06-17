from __future__ import annotations

from pathlib import Path
from xml.sax.saxutils import escape
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "report_assets"
DOCX_PATH = ROOT / "Отчет_Банковские_транзакции.docx"


def font(size: int, bold: bool = False):
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def rounded_box(draw, xy, text, fill, outline="#8EA9C1", text_fill="#102A43", width=2):
    draw.rounded_rectangle(xy, radius=12, fill=fill, outline=outline, width=width)
    x1, y1, x2, y2 = xy
    lines = text.split("\n")
    total_h = sum(draw.textbbox((0, 0), line, font=font(24, True))[3] for line in lines) + (len(lines) - 1) * 8
    y = y1 + ((y2 - y1) - total_h) / 2 - 2
    for line in lines:
        f = font(24, True)
        bbox = draw.textbbox((0, 0), line, font=f)
        draw.text((x1 + ((x2 - x1) - (bbox[2] - bbox[0])) / 2, y), line, fill=text_fill, font=f)
        y += (bbox[3] - bbox[1]) + 8


def arrow(draw, start, end, color="#46627F"):
    draw.line([start, end], fill=color, width=4)
    x1, y1 = start
    x2, y2 = end
    if x2 >= x1:
        points = [(x2, y2), (x2 - 16, y2 - 9), (x2 - 16, y2 + 9)]
    else:
        points = [(x2, y2), (x2 + 16, y2 - 9), (x2 + 16, y2 + 9)]
    draw.polygon(points, fill=color)


def draw_multiline_center(draw, box, text, size=24, bold=True, fill="#102A43"):
    x1, y1, x2, y2 = box
    lines = text.split("\n")
    fonts = [font(size, bold) for _ in lines]
    heights = [draw.textbbox((0, 0), line, font=f)[3] - draw.textbbox((0, 0), line, font=f)[1] for line, f in zip(lines, fonts)]
    total_h = sum(heights) + (len(lines) - 1) * 8
    y = y1 + ((y2 - y1) - total_h) / 2
    for line, f, h in zip(lines, fonts, heights):
        bbox = draw.textbbox((0, 0), line, font=f)
        draw.text((x1 + ((x2 - x1) - (bbox[2] - bbox[0])) / 2, y), line, fill=fill, font=f)
        y += h + 8


def idef_function_box(draw, xy, title, node, number=None):
    draw.rectangle(xy, fill="#FFFFFF", outline="#111111", width=3)
    draw_multiline_center(draw, xy, title, size=24, bold=True, fill="#111111")
    x1, y1, x2, y2 = xy
    if number is not None:
        draw.text((x1 + 12, y1 + 8), str(number), fill="#111111", font=font(20, True))
    draw.text((x2 - 58, y2 - 34), node, fill="#111111", font=font(18, True))


def label(draw, text, xy, size=18, anchor="mm"):
    draw.text(xy, text, fill="#111111", font=font(size), anchor=anchor)


def line_arrow(draw, points, color="#111111", width=3):
    draw.line(points, fill=color, width=width)
    x1, y1 = points[-2]
    x2, y2 = points[-1]
    if abs(x2 - x1) >= abs(y2 - y1):
        if x2 >= x1:
            head = [(x2, y2), (x2 - 14, y2 - 8), (x2 - 14, y2 + 8)]
        else:
            head = [(x2, y2), (x2 + 14, y2 - 8), (x2 + 14, y2 + 8)]
    else:
        if y2 >= y1:
            head = [(x2, y2), (x2 - 8, y2 - 14), (x2 + 8, y2 - 14)]
        else:
            head = [(x2, y2), (x2 - 8, y2 + 14), (x2 + 8, y2 + 14)]
    draw.polygon(head, fill=color)


def dashed_line(draw, start, end, fill="#111111", width=2, dash=10):
    x1, y1 = start
    x2, y2 = end
    length = ((x2 - x1) ** 2 + (y2 - y1) ** 2) ** 0.5
    if length == 0:
        return
    steps = int(length // dash)
    for i in range(0, steps, 2):
        a = i / steps
        b = min((i + 1) / steps, 1)
        draw.line((x1 + (x2 - x1) * a, y1 + (y2 - y1) * a, x1 + (x2 - x1) * b, y1 + (y2 - y1) * b), fill=fill, width=width)


def dfd_process(draw, xy, number, text):
    draw.ellipse(xy, fill="#FFFFFF", outline="#111111", width=3)
    x1, y1, x2, y2 = xy
    draw.text((x1 + 16, y1 + 12), number, fill="#111111", font=font(18, True))
    draw_multiline_center(draw, (x1 + 10, y1 + 35, x2 - 10, y2 - 10), text, size=18, bold=True, fill="#111111")


def dfd_store(draw, xy, code, text):
    x1, y1, x2, y2 = xy
    draw.rectangle(xy, fill="#FFFFFF", outline="#111111", width=3)
    draw.line((x1 + 55, y1, x1 + 55, y2), fill="#111111", width=3)
    draw.text((x1 + 14, y1 + 18), code, fill="#111111", font=font(18, True))
    draw_multiline_center(draw, (x1 + 60, y1, x2, y2), text, size=18, bold=True, fill="#111111")


def uml_actor(draw, center_x, top_y, name):
    draw.ellipse((center_x - 18, top_y, center_x + 18, top_y + 36), outline="#111111", width=3)
    draw.line((center_x, top_y + 36, center_x, top_y + 105), fill="#111111", width=3)
    draw.line((center_x - 45, top_y + 60, center_x + 45, top_y + 60), fill="#111111", width=3)
    draw.line((center_x, top_y + 105, center_x - 40, top_y + 155), fill="#111111", width=3)
    draw.line((center_x, top_y + 105, center_x + 40, top_y + 155), fill="#111111", width=3)
    draw_multiline_center(draw, (center_x - 80, top_y + 165, center_x + 80, top_y + 230), name, size=18, bold=True)


def uml_oval(draw, xy, text):
    draw.ellipse(xy, fill="#FFFFFF", outline="#111111", width=3)
    draw_multiline_center(draw, xy, text, size=18, bold=True, fill="#111111")


def uml_class(draw, xy, name, attributes, methods):
    x1, y1, x2, y2 = xy
    draw.rectangle(xy, fill="#FFFFFF", outline="#111111", width=3)
    header_h = 42
    attr_h = 88
    draw.line((x1, y1 + header_h, x2, y1 + header_h), fill="#111111", width=2)
    draw.line((x1, y1 + header_h + attr_h, x2, y1 + header_h + attr_h), fill="#111111", width=2)
    draw_multiline_center(draw, (x1, y1, x2, y1 + header_h), name, size=18, bold=True)
    y = y1 + header_h + 10
    for item in attributes:
        draw.text((x1 + 10, y), item, fill="#111111", font=font(14))
        y += 21
    y = y1 + header_h + attr_h + 10
    for item in methods:
        draw.text((x1 + 10, y), item, fill="#111111", font=font(14))
        y += 21


def write_drawio(name, title, vertices, edges):
    cells = [
        '<mxCell id="0"/>',
        '<mxCell id="1" parent="0"/>',
    ]
    for idx, (label_text, x, y, w, h, style) in enumerate(vertices, start=2):
        cells.append(
            f'<mxCell id="v{idx}" value="{escape(label_text)}" style="{style}" vertex="1" parent="1">'
            f'<mxGeometry x="{x}" y="{y}" width="{w}" height="{h}" as="geometry"/></mxCell>'
        )
    for idx, (label_text, points, style) in enumerate(edges, start=1):
        pts = "".join(f'<mxPoint x="{x}" y="{y}" as="{kind}"/>' for kind, x, y in points)
        cells.append(
            f'<mxCell id="e{idx}" value="{escape(label_text)}" style="{style}" edge="1" parent="1">'
            f'<mxGeometry relative="1" as="geometry">{pts}</mxGeometry></mxCell>'
        )
    xml = (
        '<mxfile host="app.diagrams.net">'
        f'<diagram name="{escape(title)}"><mxGraphModel dx="1400" dy="820" grid="1" gridSize="10" guides="1" '
        'tooltips="1" connect="1" arrows="1" fold="1" page="1" pageScale="1" '
        'pageWidth="1400" pageHeight="820" math="0" shadow="0"><root>'
        + "".join(cells)
        + '</root></mxGraphModel></diagram></mxfile>'
    )
    (OUT_DIR / name).write_text(xml, encoding="utf-8")


def save_architecture():
    img = Image.new("RGB", (1400, 760), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((50, 35), "Архитектура приложения", fill="#0B2545", font=font(36, True))
    rounded_box(draw, (70, 160, 360, 300), "WPF-клиент 1\nокно оператора", "#E8F1FB")
    rounded_box(draw, (70, 390, 360, 530), "WPF-клиент 2\nи другие окна", "#E8F1FB")
    rounded_box(draw, (540, 275, 860, 415), "TCP-сервер\nконсоль", "#F2F4F7")
    rounded_box(draw, (1030, 160, 1320, 300), "BankTransactions.Core\nбизнес-логика", "#EEF7EE")
    rounded_box(draw, (1030, 390, 1320, 530), "SQLite\nbank.db", "#FFF4D8")
    arrow(draw, (360, 230), (540, 330))
    arrow(draw, (360, 460), (540, 360))
    arrow(draw, (860, 330), (1030, 230))
    arrow(draw, (1180, 300), (1180, 390))
    draw.text((445, 210), "JSON-запросы", fill="#555555", font=font(22))
    draw.text((905, 210), "операции банка", fill="#555555", font=font(22))
    draw.text((1205, 335), "SQL", fill="#555555", font=font(22))
    path = OUT_DIR / "architecture.png"
    img.save(path)
    return path


def save_database_scheme():
    img = Image.new("RGB", (1400, 900), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((50, 35), "Схема базы данных", fill="#0B2545", font=font(36, True))
    boxes = {
        "Clients\nId, Type, Name,\nTaxNumber, Phone": (80, 160, 370, 315),
        "Accounts\nId, ClientId, Type,\nCurrency, Balance": (520, 160, 830, 315),
        "Transactions\nFromAccountId,\nToAccountId,\nAmount, Commission": (980, 160, 1320, 335),
        "PartnerBanks\nName, Country,\nIsForeign": (980, 500, 1320, 650),
        "Commissions\nAccountType,\nIsForeignPartner,\nPercent": (520, 500, 830, 650),
        "Users\nLogin, Password": (80, 500, 370, 620),
    }
    for text, xy in boxes.items():
        rounded_box(draw, xy, text, "#F7F9FC")
    arrow(draw, (370, 235), (520, 235))
    arrow(draw, (830, 235), (980, 235))
    arrow(draw, (1150, 335), (1150, 500))
    arrow(draw, (830, 575), (980, 575))
    draw.text((405, 200), "1 клиент -> N счетов", fill="#555555", font=font(20))
    draw.text((845, 200), "счета участвуют\nв транзакциях", fill="#555555", font=font(20))
    draw.text((1168, 405), "банк-партнер", fill="#555555", font=font(20))
    draw.text((842, 538), "правило комиссии", fill="#555555", font=font(20))
    path = OUT_DIR / "database_scheme.png"
    img.save(path)
    return path


def save_sequence():
    img = Image.new("RGB", (1400, 820), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((50, 35), "UML: диаграмма последовательности", fill="#0B2545", font=font(36, True))
    columns = [160, 470, 790, 1110]
    names = ["Оператор", "Клиентское окно", "Сервер", "SQLite"]
    for x, name in zip(columns, names):
        draw.rectangle((x - 105, 135, x + 105, 190), fill="#FFFFFF", outline="#111111", width=3)
        draw_multiline_center(draw, (x - 105, 135, x + 105, 190), name, size=18, bold=True)
        dashed_line(draw, (x, 190), (x, 735), fill="#777777", width=2, dash=12)
    for x, y1, y2 in [(470, 245, 675), (790, 270, 700), (1110, 320, 610)]:
        draw.rectangle((x - 8, y1, x + 8, y2), fill="#FFFFFF", outline="#111111", width=2)
    messages = [
        (235, 160, 470, "1: ввод логина и пароля", False),
        (275, 470, 790, "2: Login(login, password)", False),
        (325, 790, 1110, "3: SELECT Users", False),
        (375, 1110, 790, "4: token", True),
        (450, 160, 470, "5: ввод данных перевода", False),
        (500, 470, 790, "6: Transfer(from, to, amount)", False),
        (555, 790, 1110, "7: UPDATE / INSERT", False),
        (625, 1110, 790, "8: ok / error", True),
        (675, 790, 470, "9: результат операции", True),
    ]
    for y, x1, x2, text, dashed in messages:
        if dashed:
            dashed_line(draw, (x1, y), (x2, y), fill="#111111", width=2, dash=12)
            line_arrow(draw, [(x1, y), (x2, y)], width=0)
        else:
            line_arrow(draw, [(x1, y), (x2, y)], width=2)
        draw.text((min(x1, x2) + 18, y - 24), text, fill="#111111", font=font(16))
    write_drawio(
        "sequence.drawio",
        "UML Sequence",
        [(name, x - 105, 135, 210, 55, "whiteSpace=wrap;html=1;rounded=0;strokeColor=#111111;fillColor=#ffffff;strokeWidth=2;") for x, name in zip(columns, names)],
        [],
    )
    path = OUT_DIR / "sequence.png"
    img.save(path)
    return path


def save_simple_flow(name, title, nodes, links):
    img = Image.new("RGB", (1400, 820), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((50, 35), title, fill="#0B2545", font=font(36, True))
    positions = {}
    for key, text, xy, fill in nodes:
        positions[key] = xy
        rounded_box(draw, xy, text, fill)
    for start_key, end_key, label in links:
        sx1, sy1, sx2, sy2 = positions[start_key]
        ex1, ey1, ex2, ey2 = positions[end_key]
        start = (sx2, (sy1 + sy2) // 2) if sx2 < ex1 else ((sx1 + sx2) // 2, sy2)
        end = (ex1, (ey1 + ey2) // 2) if sx2 < ex1 else ((ex1 + ex2) // 2, ey1)
        arrow(draw, start, end)
        if label:
            lx = (start[0] + end[0]) // 2 - 90
            ly = (start[1] + end[1]) // 2 - 30
            draw.text((lx, ly), label, fill="#555555", font=font(18))
    path = OUT_DIR / name
    img.save(path)
    return path


def save_idef0_context():
    img = Image.new("RGB", (1400, 820), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((50, 35), "IDEF0 A-0: контекстная диаграмма", fill="#0B2545", font=font(36, True))
    draw.rectangle((45, 105, 1355, 755), outline="#111111", width=2)
    box = (480, 310, 920, 520)
    idef_function_box(draw, box, "Управлять\nбанковскими\nтранзакциями", "A-0")

    line_arrow(draw, [(90, 390), (480, 390)])
    label(draw, "Данные клиента", (245, 360))
    line_arrow(draw, [(90, 450), (480, 450)])
    label(draw, "Данные операции", (250, 420))

    line_arrow(draw, [(570, 145), (570, 310)])
    label(draw, "Правила банка", (570, 125), size=16)
    line_arrow(draw, [(700, 145), (700, 310)])
    label(draw, "Комиссии", (700, 125), size=16)
    line_arrow(draw, [(830, 145), (830, 310)])
    label(draw, "Авторизация", (830, 125), size=16)

    line_arrow(draw, [(920, 390), (1310, 390)])
    label(draw, "Обновленные счета", (1095, 360))
    line_arrow(draw, [(920, 450), (1310, 450)])
    label(draw, "История транзакций", (1100, 420))

    line_arrow(draw, [(570, 710), (570, 520)])
    label(draw, "Оператор", (570, 735))
    line_arrow(draw, [(700, 710), (700, 520)])
    label(draw, "C#/.NET", (700, 735))
    line_arrow(draw, [(830, 710), (830, 520)])
    label(draw, "SQLite", (830, 735))

    draw.text((65, 715), "NODE: A-0", fill="#111111", font=font(18, True))
    draw.text((975, 715), "TITLE: Банковские транзакции", fill="#111111", font=font(18, True))
    write_drawio(
        "idef0_context.drawio",
        "IDEF0 A-0",
        [("Управлять<br>банковскими<br>транзакциями<br><br>A-0", 520, 310, 360, 210, "whiteSpace=wrap;html=1;rounded=0;strokeColor=#111111;fillColor=#ffffff;strokeWidth=3;")],
        [],
    )
    path = OUT_DIR / "idef0_context.png"
    img.save(path)
    return path


def save_idef0_decomposition():
    img = Image.new("RGB", (1400, 900), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((50, 35), "IDEF0 A0: декомпозиция процесса", fill="#0B2545", font=font(36, True))
    draw.rectangle((45, 105, 1355, 825), outline="#111111", width=2)
    boxes = {
        "a1": (120, 210, 380, 340, "Авторизовать\nоператора", "A1", 1),
        "a2": (455, 250, 715, 380, "Вести клиентов\nи счета", "A2", 2),
        "a3": (790, 290, 1050, 420, "Рассчитать\nкомиссию", "A3", 3),
        "a4": (455, 530, 715, 660, "Провести\nтранзакцию", "A4", 4),
        "a5": (790, 570, 1050, 700, "Сохранить\nисторию", "A5", 5),
    }
    for _, (x1, y1, x2, y2, text, node, number) in boxes.items():
        idef_function_box(draw, (x1, y1, x2, y2), text, node, number)

    line_arrow(draw, [(75, 275), (120, 275)])
    label(draw, "Логин,\nпароль", (90, 235), size=17)
    line_arrow(draw, [(75, 610), (455, 610)])
    label(draw, "Данные клиента,\nсчета и операции", (250, 575), size=17)

    line_arrow(draw, [(250, 145), (250, 210)])
    label(draw, "Политика доступа", (250, 125), size=17)
    line_arrow(draw, [(585, 145), (585, 250)])
    label(draw, "Правила учета", (585, 125), size=17)
    line_arrow(draw, [(920, 145), (920, 290)])
    label(draw, "Тарифы комиссий", (920, 125), size=17)

    line_arrow(draw, [(380, 275), (455, 315)])
    label(draw, "токен", (420, 275), size=17)
    line_arrow(draw, [(715, 315), (790, 355)])
    label(draw, "тип счета\nи банк", (750, 300), size=17)
    line_arrow(draw, [(920, 420), (920, 500), (715, 500), (715, 590)])
    label(draw, "комиссия", (825, 475), size=17)
    line_arrow(draw, [(715, 610), (790, 635)])
    label(draw, "операция", (755, 605), size=17)

    line_arrow(draw, [(715, 545), (1315, 545)])
    label(draw, "Обновленные счета", (1080, 515), size=17)
    line_arrow(draw, [(1050, 635), (1315, 635)])
    label(draw, "История транзакций", (1180, 605), size=17)

    line_arrow(draw, [(585, 785), (585, 660)])
    label(draw, "Сервер C#", (585, 808), size=17)
    line_arrow(draw, [(920, 785), (920, 700)])
    label(draw, "SQLite", (920, 808), size=17)
    line_arrow(draw, [(250, 785), (250, 340)])
    label(draw, "Оператор", (250, 808), size=17)

    draw.text((65, 790), "NODE: A0", fill="#111111", font=font(18, True))
    draw.text((970, 790), "TITLE: Управлять банковскими транзакциями", fill="#111111", font=font(18, True))
    write_drawio(
        "idef0_decomposition.drawio",
        "IDEF0 A0",
        [(data[4] + f"<br><br>{data[5]}", data[0], data[1], data[2] - data[0], data[3] - data[1], "whiteSpace=wrap;html=1;rounded=0;strokeColor=#111111;fillColor=#ffffff;strokeWidth=3;") for data in boxes.values()],
        [],
    )
    path = OUT_DIR / "idef0_decomposition.png"
    img.save(path)
    return path


def save_idef3():
    img = Image.new("RGB", (1400, 820), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((50, 35), "IDEF3: сценарий проведения перевода", fill="#0B2545", font=font(36, True))
    draw.rectangle((45, 105, 1355, 755), outline="#111111", width=2)
    uobs = [
        ("UOB 1", "Выполнить\nавторизацию", (90, 300, 300, 430)),
        ("UOB 2", "Выбрать счет\nсписания и счет\nполучателя", (360, 300, 610, 430)),
        ("UOB 3", "Проверить баланс\nс учетом комиссии", (670, 300, 920, 430)),
        ("UOB 4", "Списать сумму\nи комиссию", (1010, 210, 1260, 340)),
        ("UOB 5", "Записать отказ\nв операции", (690, 520, 940, 650)),
        ("UOB 6", "Сохранить\nтранзакцию", (1010, 610, 1260, 710)),
    ]
    for code, text, xy in uobs:
        draw.rectangle(xy, fill="#FFFFFF", outline="#111111", width=3)
        x1, y1, x2, y2 = xy
        draw.text((x1 + 8, y1 + 6), code, fill="#111111", font=font(16, True))
        draw_multiline_center(draw, (x1, y1 + 22, x2, y2), text, size=20, bold=True, fill="#111111")
    junction = (965, 365)
    draw.ellipse((junction[0] - 22, junction[1] - 22, junction[0] + 22, junction[1] + 22), fill="#FFFFFF", outline="#111111", width=3)
    draw.text((junction[0] - 6, junction[1] - 14), "X", fill="#111111", font=font(22, True))
    line_arrow(draw, [(300, 365), (360, 365)])
    line_arrow(draw, [(610, 365), (670, 365)])
    line_arrow(draw, [(920, 365), (943, 365)])
    line_arrow(draw, [(987, 365), (1010, 275)])
    label(draw, "достаточно средств", (1110, 190), size=16)
    line_arrow(draw, [(965, 387), (965, 585), (940, 585)])
    label(draw, "недостаточно средств", (850, 495), size=16)
    line_arrow(draw, [(1135, 340), (1135, 610)])
    label(draw, "после списания", (1225, 480), size=16)
    draw.text((65, 710), "SCENARIO: перевод между счетами", fill="#111111", font=font(18, True))
    draw.text((1135, 710), "IDEF3: UOB + XOR junction", fill="#111111", font=font(18, True))
    write_drawio(
        "idef3.drawio",
        "IDEF3",
        [(f"{code}<br>{text.replace(chr(10), '<br>')}", xy[0], xy[1], xy[2] - xy[0], xy[3] - xy[1], "whiteSpace=wrap;html=1;rounded=0;strokeColor=#111111;fillColor=#ffffff;strokeWidth=3;") for code, text, xy in uobs],
        [],
    )
    path = OUT_DIR / "idef3.png"
    img.save(path)
    return path


def save_dfd():
    img = Image.new("RGB", (1400, 900), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((50, 35), "DFD Level 0: диаграмма потоков данных", fill="#0B2545", font=font(36, True))

    draw.rectangle((70, 385, 270, 495), fill="#FFFFFF", outline="#111111", width=3)
    draw_multiline_center(draw, (70, 385, 270, 495), "Оператор\nбанка", size=18, bold=True)

    dfd_process(draw, (380, 130, 600, 310), "1.0", "Проверить\nавторизацию")
    dfd_process(draw, (380, 390, 600, 570), "2.0", "Вести\nклиентов\nи счета")
    dfd_process(draw, (780, 390, 1000, 570), "3.0", "Провести\nтранзакцию")
    dfd_process(draw, (1090, 130, 1310, 310), "4.0", "Сформировать\nисторию")

    dfd_store(draw, (720, 160, 1000, 240), "D1", "Users")
    dfd_store(draw, (260, 700, 540, 780), "D2", "Clients")
    dfd_store(draw, (620, 700, 900, 780), "D3", "Accounts")
    dfd_store(draw, (1030, 700, 1310, 780), "D4", "Transactions")
    dfd_store(draw, (1030, 520, 1310, 600), "D5", "Commissions,\nPartnerBanks")

    flows = [
        ([(270, 405), (380, 215)], "логин, пароль", (300, 285)),
        ([(600, 190), (720, 190)], "запрос", (640, 165)),
        ([(720, 220), (600, 245)], "пользователь", (625, 245)),
        ([(270, 455), (380, 480)], "клиент / счет", (285, 465)),
        ([(490, 570), (400, 700)], "данные клиента", (410, 635)),
        ([(545, 570), (690, 700)], "данные счета", (570, 635)),
        ([(600, 480), (780, 480)], "операция", (665, 450)),
        ([(760, 700), (840, 570)], "счета", (765, 620)),
        ([(1000, 445), (1030, 560)], "тарифы", (1010, 495)),
        ([(890, 570), (890, 700)], "изменения", (905, 625)),
        ([(1000, 480), (1090, 220)], "результат", (1020, 345)),
        ([(1200, 310), (1170, 700)], "запись", (1210, 500)),
        ([(1090, 220), (1090, 95), (170, 95), (170, 385)], "ответ", (520, 105)),
    ]
    for points, text, text_xy in flows:
        line_arrow(draw, points, width=2)
        draw.text(text_xy, text, fill="#111111", font=font(14))

    write_drawio(
        "dfd.drawio",
        "DFD Level 0",
        [
            ("Оператор<br>банка", 70, 370, 200, 110, "whiteSpace=wrap;html=1;rounded=0;strokeColor=#111111;fillColor=#ffffff;strokeWidth=2;"),
            ("1.0<br>Проверить<br>авторизацию", 390, 140, 220, 190, "ellipse;whiteSpace=wrap;html=1;strokeColor=#111111;fillColor=#ffffff;strokeWidth=2;"),
            ("2.0<br>Вести<br>клиентов<br>и счета", 390, 360, 220, 190, "ellipse;whiteSpace=wrap;html=1;strokeColor=#111111;fillColor=#ffffff;strokeWidth=2;"),
            ("3.0<br>Провести<br>транзакцию", 725, 360, 220, 190, "ellipse;whiteSpace=wrap;html=1;strokeColor=#111111;fillColor=#ffffff;strokeWidth=2;"),
            ("4.0<br>Сформировать<br>историю", 1050, 360, 220, 190, "ellipse;whiteSpace=wrap;html=1;strokeColor=#111111;fillColor=#ffffff;strokeWidth=2;"),
        ],
        [],
    )
    path = OUT_DIR / "dfd.png"
    img.save(path)
    return path


def save_usecase():
    img = Image.new("RGB", (1400, 850), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((50, 35), "UML: диаграмма вариантов использования", fill="#0B2545", font=font(36, True))
    uml_actor(draw, 165, 320, "Оператор\nбанка")
    draw.rectangle((350, 125, 1290, 760), fill="#FFFFFF", outline="#111111", width=3)
    draw.text((375, 145), "Система банковских транзакций", fill="#111111", font=font(20, True))
    cases = {
        "login": ((570, 175, 930, 255), "Войти\nв систему"),
        "client": ((570, 270, 930, 350), "Управлять\nклиентами"),
        "account": ((570, 365, 930, 445), "Управлять\nсчетами"),
        "deposit": ((570, 460, 930, 540), "Пополнить\nсчет"),
        "transfer": ((570, 555, 930, 635), "Выполнить\nперевод"),
        "delete": ((570, 650, 930, 730), "Удалить\nклиента"),
    }
    for xy, text in cases.values():
        uml_oval(draw, xy, text)
    association_points = {
        "login": (245, 345),
        "client": (245, 365),
        "account": (245, 385),
        "deposit": (245, 405),
        "transfer": (245, 425),
        "delete": (245, 445),
    }
    for key in ["login", "client", "account", "deposit", "transfer", "delete"]:
        x1, y1, x2, y2 = cases[key][0]
        draw.line((association_points[key], (x1, (y1 + y2) // 2)), fill="#111111", width=2)
    x1, y1, x2, y2 = cases["delete"][0]
    lx1, ly1, lx2, ly2 = cases["login"][0]
    dashed_line(draw, (x2, (y1 + y2) // 2), (x2 + 105, (y1 + y2) // 2), width=2, dash=12)
    dashed_line(draw, (x2 + 105, (y1 + y2) // 2), (x2 + 105, (ly1 + ly2) // 2), width=2, dash=12)
    dashed_line(draw, (x2 + 105, (ly1 + ly2) // 2), (lx2, (ly1 + ly2) // 2), width=2, dash=12)
    draw.text((960, 440), "<<include>>", fill="#111111", font=font(13))
    write_drawio(
        "usecase.drawio",
        "UML Use Case",
        [("Оператор<br>банка", 95, 320, 140, 230, "shape=umlActor;verticalLabelPosition=bottom;verticalAlign=top;html=1;outlineConnect=0;strokeColor=#111111;"),
         ("Система банковских транзакций", 350, 125, 940, 635, "whiteSpace=wrap;html=1;rounded=0;strokeColor=#111111;fillColor=none;strokeWidth=2;")]
        + [(text.replace("\n", "<br>"), xy[0], xy[1], xy[2] - xy[0], xy[3] - xy[1], "ellipse;whiteSpace=wrap;html=1;strokeColor=#111111;fillColor=#ffffff;strokeWidth=2;") for xy, text in cases.values()],
        [],
    )
    path = OUT_DIR / "usecase.png"
    img.save(path)
    return path


def save_class_diagram():
    img = Image.new("RGB", (1500, 1060), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((50, 35), "UML: диаграмма классов", fill="#0B2545", font=font(36, True))
    classes = {
        "client": ((70, 135, 360, 335), "Client", ["+ Id: int", "+ Type: ClientType", "+ Name: string", "+ TaxNumber: string"], []),
        "account": ((440, 135, 760, 335), "Account", ["+ Id: int", "+ ClientId: int", "+ Type: AccountType", "+ Balance: decimal"], []),
        "tx": ((850, 135, 1245, 355), "BankTransaction", ["+ Id: int", "+ Type: TransactionType", "+ Amount: decimal", "+ Commission: decimal"], []),
        "partner": ((850, 450, 1245, 650), "PartnerBank", ["+ Id: int", "+ Name: string", "+ Country: string", "+ IsForeign: bool"], []),
        "commission": ((440, 450, 760, 650), "CommissionRule", ["+ Id: int", "+ AccountType: AccountType", "+ IsForeignPartner: bool", "+ Percent: decimal"], []),
        "service": ((70, 735, 470, 1005), "BankService", ["- database: BankDatabase"], ["+ Login()", "+ CreateClient()", "+ Deposit()", "+ Transfer()"]),
        "db": ((560, 735, 930, 1005), "BankDatabase", ["- connectionString: string"], ["+ Initialize()", "+ OpenConnection()"]),
    }
    for xy, name, attrs, methods in classes.values():
        uml_class(draw, xy, name, attrs, methods)
    line_arrow(draw, [(360, 235), (440, 235)], width=2)
    draw.text((370, 205), "1", fill="#111111", font=font(15))
    draw.text((405, 205), "0..*", fill="#111111", font=font(15))
    draw.text((385, 245), "имеет", fill="#111111", font=font(14))
    line_arrow(draw, [(760, 235), (850, 235)], width=2)
    draw.text((770, 205), "1", fill="#111111", font=font(15))
    draw.text((812, 205), "0..*", fill="#111111", font=font(15))
    draw.text((775, 245), "участвует", fill="#111111", font=font(14))
    line_arrow(draw, [(1045, 355), (1045, 450)], width=2)
    draw.text((1060, 395), "0..1", fill="#111111", font=font(15))
    line_arrow(draw, [(760, 550), (850, 550)], width=2)
    draw.text((775, 520), "используется", fill="#111111", font=font(14))
    dashed_line(draw, (470, 855), (560, 855), width=2, dash=12)
    line_arrow(draw, [(470, 855), (560, 855)], width=0)
    draw.text((485, 825), "<<uses>>", fill="#111111", font=font(14))
    write_drawio(
        "class_diagram.drawio",
        "UML Class Diagram",
        [(name + "<hr/>" + "<br>".join(attrs) + "<hr/>" + "<br>".join(methods), xy[0], xy[1], xy[2] - xy[0], xy[3] - xy[1], "swimlane;fontStyle=1;align=center;verticalAlign=top;childLayout=stackLayout;horizontal=1;startSize=32;horizontalStack=0;resizeParent=1;resizeParentMax=0;resizeLast=0;collapsible=0;marginBottom=0;whiteSpace=wrap;html=1;strokeColor=#111111;fillColor=#ffffff;strokeWidth=2;") for xy, name, attrs, methods in classes.values()],
        [],
    )
    path = OUT_DIR / "class_diagram.png"
    img.save(path)
    return path


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, "F2F4F7")
        set_cell_width(cell, widths[idx])
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            set_cell_width(cells[idx], widths[idx])
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc.add_paragraph()
    return table


def add_code_block(doc, title, code, explanation):
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor(31, 77, 120)
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "F7F7F7")
    set_cell_width(cell, 9360)
    cell.text = ""
    p = cell.paragraphs[0]
    for line in code.strip().splitlines():
        run = p.add_run(line.rstrip())
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        run.add_break()
    p = doc.add_paragraph(explanation)
    p.paragraph_format.space_after = Pt(8)


def add_image_with_caption(doc, image_path, caption):
    doc.add_picture(str(image_path), width=Inches(6.2))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(85, 85, 85)


def setup_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, RGBColor(46, 116, 181), 16, 8),
        ("Heading 2", 13, RGBColor(46, 116, 181), 12, 6),
        ("Heading 3", 12, RGBColor(31, 77, 120), 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Банковские транзакции")


def build_doc():
    OUT_DIR.mkdir(exist_ok=True)
    architecture = save_architecture()
    database = save_database_scheme()
    sequence = save_sequence()
    idef0_context = save_idef0_context()
    idef0_decomposition = save_idef0_decomposition()
    idef3 = save_idef3()
    dfd = save_dfd()
    usecase = save_usecase()
    class_diagram = save_class_diagram()

    doc = Document()
    setup_styles(doc)
    add_footer(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Курсовой проект\nпо дисциплине\n«Методы и средства проектирования информационных систем и технологий»\n\nНа тему\n«Банковские транзакции»")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(11, 37, 69)
    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    meta.add_run("Выполнил: ______________________\nПроверил: ______________________")
    doc.add_paragraph()
    city = doc.add_paragraph("2026")
    city.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    doc.add_heading("Содержание", level=1)
    for item in [
        "1. Задание на курсовой проект",
        "2. Предметная область и архитектура системы",
        "3. Проектирование системы",
        "4. Схема базы данных",
        "5. Реализация взаимодействия с БД",
        "6. Авторизация и доступ к операциям",
        "7. Тестирование",
        "8. Порядок запуска exe",
        "9. Заключение",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_page_break()

    doc.add_heading("1. Задание на курсовой проект", level=1)
    doc.add_paragraph(
        "Необходимо разработать систему регистрации банковских транзакций и обновления "
        "счетов клиентов. В системе должны поддерживаться клиенты-физические лица и "
        "клиенты-юридические лица, несколько счетов у одного клиента, разные типы счетов, "
        "банки-партнеры, комиссии и история операций."
    )
    for item in [
        "База данных должна содержать данные о клиентах, счетах, транзакциях, комиссиях и банках-партнерах.",
        "Приложение должно состоять из базы данных, серверной консольной части и оконной клиентской части.",
        "Клиентских окон может быть запущено несколько одновременно.",
        "Сервер должен отслеживать подключенных клиентов.",
        "Для удаления клиента необходимо предварительно выполнить вход в систему.",
        "В отчет должны входить unit-тесты и сценарий демонстрации.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("2. Предметная область и архитектура системы", level=1)
    doc.add_heading("2.1. Предметная область", level=2)
    doc.add_paragraph(
        "Предметная область связана с банковским обслуживанием клиентов. Оператор банка "
        "создает карточки клиентов, открывает счета, пополняет баланс и выполняет переводы. "
        "При переводе система рассчитывает комиссию в зависимости от типа счета и банка-партнера."
    )
    doc.add_heading("2.2. Технологический стек", level=2)
    add_table(
        doc,
        ["Компонент", "Выбор", "Назначение"],
        [
            ["Язык", "C#", "Единый язык для серверной, оконной и тестовой частей."],
            ["База данных", "SQLite", "Локальный файл БД без отдельного сервера."],
            ["Сервер", "Console App + TCP", "Принимает запросы от нескольких клиентов и пишет в БД."],
            ["Клиент", "WPF", "Оконный интерфейс оператора банка."],
            ["Тесты", "xUnit", "Проверка входа, пополнения, перевода и ошибок баланса."],
        ],
        [1900, 2100, 5360],
    )
    doc.add_heading("2.3. Структура проекта", level=2)
    add_table(
        doc,
        ["Проект", "Назначение"],
        [
            ["BankTransactions.Core", "Общие модели, создание SQLite-базы и бизнес-логика операций."],
            ["BankTransactions.Server", "Консольный TCP-сервер, принимающий запросы клиентов."],
            ["BankTransactions.Client", "WPF-окно оператора банка."],
            ["BankTransactions.Tests", "Unit-тесты бизнес-логики."],
            ["dist", "Готовые exe-файлы и bat-запускатели."],
        ],
        [2900, 6460],
    )
    doc.add_heading("2.4. Взаимодействие компонентов", level=2)
    doc.add_paragraph(
        "Приложение разделено на общую библиотеку, сервер и клиент. Клиентские окна не "
        "работают с базой напрямую: они отправляют JSON-команды на сервер. Сервер хранит "
        "состояние подключений и выполняет операции через библиотеку BankTransactions.Core."
    )
    add_image_with_caption(doc, architecture, "Рисунок 1 - Архитектура приложения")

    doc.add_heading("3. Проектирование системы", level=1)
    doc.add_heading("3.1. IDEF0 A-0 - контекстная диаграмма", level=2)
    add_image_with_caption(doc, idef0_context, "Рисунок 2 - Контекстная диаграмма IDEF0")
    doc.add_heading("3.2. IDEF0 A0 - декомпозиция", level=2)
    add_image_with_caption(doc, idef0_decomposition, "Рисунок 3 - Декомпозиция основного процесса")
    doc.add_heading("3.3. IDEF3 - сценарий процесса", level=2)
    add_image_with_caption(doc, idef3, "Рисунок 4 - Сценарий проведения перевода")
    doc.add_heading("3.4. DFD - диаграмма потоков данных", level=2)
    add_image_with_caption(doc, dfd, "Рисунок 5 - Потоки данных между оператором, клиентом, сервером и БД")
    doc.add_heading("3.5. UML - диаграмма вариантов использования", level=2)
    add_image_with_caption(doc, usecase, "Рисунок 6 - Варианты использования системы")
    doc.add_heading("3.6. UML - диаграмма классов", level=2)
    add_image_with_caption(doc, class_diagram, "Рисунок 7 - Основные классы предметной области")
    doc.add_heading("3.7. UML - диаграмма последовательности", level=2)
    add_image_with_caption(doc, sequence, "Рисунок 8 - Последовательность выполнения перевода")

    doc.add_heading("4. Схема базы данных", level=1)
    doc.add_paragraph(
        "База данных создается автоматически при первом запуске сервера. Для учебного проекта "
        "используется SQLite: она хранит данные в одном локальном файле bank.db и не требует "
        "отдельной установки СУБД."
    )
    add_image_with_caption(doc, database, "Рисунок 9 - Основные таблицы и связи")
    add_table(
        doc,
        ["Таблица", "Назначение"],
        [
            ["Users", "Хранит учетные данные для входа. Начальная учетная запись: admin/admin."],
            ["Clients", "Клиенты банка: физические и юридические лица."],
            ["Accounts", "Счета клиентов: зарплатные, валютные и накопительные."],
            ["PartnerBanks", "Банки-партнеры, включая зарубежные."],
            ["Commissions", "Проценты комиссий по типу счета и признаку зарубежного партнера."],
            ["Transactions", "История пополнений и переводов с суммой и комиссией."],
        ],
        [2200, 7160],
    )

    doc.add_heading("5. Реализация взаимодействия с БД", level=1)
    add_code_block(
        doc,
        "Создание таблиц SQLite",
        """
CREATE TABLE IF NOT EXISTS Clients (
    Id INTEGER PRIMARY KEY AUTOINCREMENT,
    Type INTEGER NOT NULL,
    Name TEXT NOT NULL,
    TaxNumber TEXT NOT NULL,
    Phone TEXT NOT NULL
);

CREATE TABLE IF NOT EXISTS Accounts (
    Id INTEGER PRIMARY KEY AUTOINCREMENT,
    ClientId INTEGER NOT NULL,
    Type INTEGER NOT NULL,
    Currency TEXT NOT NULL,
    Balance TEXT NOT NULL,
    FOREIGN KEY (ClientId) REFERENCES Clients(Id) ON DELETE CASCADE
);
        """,
        "Фрагмент показывает базовые сущности системы: клиент может иметь несколько счетов, а счета связаны с клиентами внешним ключом.",
    )
    add_code_block(
        doc,
        "Расчет комиссии и перевод",
        """
var commission = CalculateCommission(connection, from.Type, partner.IsForeign, request.Amount);
var total = request.Amount + commission;

if (from.Balance < total)
{
    throw new InvalidOperationException("Недостаточно средств с учетом комиссии.");
}

SetBalance(connection, request.FromAccountId, from.Balance - total);
SetBalance(connection, request.ToAccountId, toBalance + request.Amount);
        """,
        "Сначала рассчитывается комиссия. Затем сервер проверяет, хватает ли денег на сумму перевода и комиссию, после чего обновляет оба счета.",
    )

    doc.add_heading("6. Авторизация и доступ к операциям", level=1)
    doc.add_paragraph(
        "В системе предусмотрен вход оператора. Начальная учетная запись создается при "
        "инициализации базы данных: логин admin, пароль admin. Операции изменения данных "
        "требуют токен авторизованной сессии."
    )
    add_code_block(
        doc,
        "Отслеживание подключенных клиентов",
        """
var sessions = new ConcurrentDictionary<string, TcpClient>();

var client = await listener.AcceptTcpClientAsync();
var sessionId = Guid.NewGuid().ToString("N");
sessions[sessionId] = client;

Console.WriteLine($"Подключенных клиентов: {sessions.Count}");
        """,
        "Консольный сервер хранит активные подключения в потокобезопасной коллекции и выводит количество подключенных окон-клиентов.",
    )

    doc.add_heading("7. Тестирование", level=1)
    doc.add_heading("7.1. Автоматические unit-тесты", level=2)
    add_code_block(
        doc,
        "Unit-тест перевода",
        """
service.Deposit(new DepositRequest(fromAccountId, 1000m, "Начальный баланс"));
service.Transfer(new TransferRequest(fromAccountId, toAccountId, 2, 100m, "Международный перевод"));

Assert.Equal(898.50m, from.Balance);
Assert.Equal(100m, to.Balance);
Assert.Equal(1.50m, transfer.Commission);
        """,
        "Тест фиксирует ожидаемое поведение: со счета отправителя списывается сумма и комиссия, а получателю зачисляется сумма перевода.",
    )
    add_table(
        doc,
        ["Тест", "Что проверяет", "Результат"],
        [
            ["Login_ReturnsTrue_ForDefaultAdmin", "Вход под учетной записью admin/admin.", "Пройден"],
            ["Deposit_IncreasesAccountBalance", "Пополнение увеличивает баланс счета.", "Пройден"],
            ["Transfer_WithdrawsAmountAndCommission", "Перевод списывает сумму и комиссию.", "Пройден"],
            ["Transfer_Throws_WhenBalanceIsNotEnough", "Недостаточный баланс вызывает ошибку.", "Пройден"],
        ],
        [2900, 4860, 1600],
    )
    doc.add_paragraph("Последний запуск: dotnet test. Результат: 4 теста пройдены, 0 ошибок.")
    doc.add_heading("7.2. Ручной сквозной сценарий", level=2)
    for item in [
        "Запустить сервер и убедиться, что консоль ожидает подключения.",
        "Запустить два окна клиента и проверить изменение количества подключенных клиентов в консоли.",
        "Войти в одном из окон под admin/admin.",
        "Создать физическое и юридическое лицо.",
        "Создать зарплатный и валютный счета.",
        "Пополнить счет отправителя.",
        "Выполнить перевод через зарубежный банк-партнер.",
        "Проверить запись в истории транзакций и рассчитанную комиссию.",
    ]:
        doc.add_paragraph(item, style="List Number")

    doc.add_heading("8. Порядок запуска exe", level=1)
    for item in [
        "Открыть папку dist.",
        "Запустить файл «Запустить сервер.bat» или dist\\Server\\BankTransactions.Server.exe.",
        "Запустить файл «Запустить клиент.bat» или dist\\Client\\BankTransactions.Client.exe.",
        "Войти в клиенте: логин admin, пароль admin.",
    ]:
        doc.add_paragraph(item, style="List Number")

    doc.add_heading("9. Вывод", level=1)
    doc.add_paragraph(
        "Разработана система банковских транзакций с локальной базой данных SQLite, "
        "консольным сервером, оконным клиентом и набором unit-тестов. Система позволяет "
        "создавать клиентов и счета, выполнять пополнения и переводы, рассчитывать комиссии "
        "и хранить историю операций."
    )

    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    print(build_doc())
